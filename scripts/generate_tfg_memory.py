from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Memoria_TFG_MyJornia.docx"
ICON = ROOT / "assets" / "icon.png"


ACCENT = RGBColor(37, 99, 235)
MUTED = RGBColor(71, 85, 105)
DARK = RGBColor(15, 23, 42)
LIGHT_FILL = "EFF6FF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def setup_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, ACCENT),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color

    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    return doc


def add_plain_paragraph(doc: Document, text: str, *, no_indent: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if no_indent:
        p.paragraph_format.first_line_indent = Cm(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], LIGHT_FILL)
        for paragraph in header_cells[idx].paragraphs:
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def add_figure_placeholder(doc: Document, number: int, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Insertar captura: {title}]")
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_caption(doc, f"Figura {number}. {title}. Fuente: elaboración propia.")


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    if ICON.exists():
        p.add_run().add_picture(str(ICON), width=Cm(3.2))

    title = doc.add_paragraph("MyJornia: aplicación móvil para la gestión de turnos y cálculo de nómina")
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)

    subtitle = doc.add_paragraph("Trabajo Final de Grado")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Cm(0)
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.color.rgb = DARK

    data = [
        ("Autor", "Jairo Panait García"),
        ("Centro", "Medac Davante Formacciona"),
        ("Ciclo", "Desarrollo de Aplicaciones Multiplataforma"),
        ("Tutor/a", "Paula Girones Gracia"),
        ("Tema", "Aplicación móvil de control de turnos y nóminas"),
        ("Fecha de entrega", "11 de mayo de 2026"),
    ]
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(data):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value
        set_cell_shading(table.rows[idx].cells[0], LIGHT_FILL)
        for cell in table.rows[idx].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            for run in table.rows[idx].cells[0].paragraphs[0].runs:
                run.bold = True

    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.paragraph_format.first_line_indent = Cm(0)
    repo.add_run("Repositorio GitHub: ").bold = True
    repo.add_run("https://github.com/jairopanait/MyJornia")

    db = doc.add_paragraph()
    db.alignment = WD_ALIGN_PARAGRAPH.CENTER
    db.paragraph_format.first_line_indent = Cm(0)
    db.add_run("Base de datos: ").bold = True
    db.add_run("Supabase PostgreSQL. URL pública del servicio: https://rhpvhzrmwzorhaqiwhhm.supabase.co. No se incluyen claves ni credenciales.")

    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Índice", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_toc(p)
    note = doc.add_paragraph("Nota: al pegar el contenido en la plantilla oficial, actualizar el índice desde Word para ajustar la paginación final.")
    note.paragraph_format.first_line_indent = Cm(0)
    note.runs[0].italic = True
    note.runs[0].font.color.rgb = MUTED
    doc.add_page_break()


def add_resumen(doc: Document) -> None:
    doc.add_heading("1. Resumen", level=1)
    add_plain_paragraph(
        doc,
        "El presente Trabajo Final de Grado aborda el desarrollo de MyJornia, una aplicación móvil orientada a trabajadores que realizan turnos variables y necesitan controlar de manera más clara su calendario laboral, las horas trabajadas y la estimación económica asociada a su nómina. El problema tecnológico detectado se relaciona con la dispersión de información entre calendarios, notas personales y documentos de empresa, lo que dificulta comprobar con precisión horas ordinarias, complementarias, nocturnas, festivas y conceptos salariales personalizados.",
    )
    add_plain_paragraph(
        doc,
        "El objetivo principal del sistema consiste en desarrollar una herramienta multiplataforma que permita registrar turnos, configurar reglas de nómina y consultar resúmenes mensuales de forma visual. Para ello se han definido objetivos específicos vinculados al diseño de una interfaz móvil intuitiva, la gestión persistente de datos, la autenticación de usuarios y la incorporación de medidas de seguridad adecuadas para información laboral sensible.",
    )
    add_plain_paragraph(
        doc,
        "La metodología utilizada se basa en un MVP con prototipado incremental durante un periodo aproximado de un mes. El desarrollo se realizó con React Native, Expo, TypeScript y Supabase, apoyándose en GitHub para el control de versiones. Se llevaron a cabo fases de análisis, diseño, implementación, pruebas funcionales, pruebas de usabilidad y validación en Expo Go, navegador web y usuarios de prueba.",
    )
    add_plain_paragraph(
        doc,
        "Como resultado se obtuvo una aplicación funcional con registro e inicio de sesión, calendario de turnos, plantillas personalizables, configuración de nómina, cálculo de nocturnidad, festivos, deducciones, perfil y opciones de seguridad. La solución mejora la visibilidad del trabajo realizado y facilita al usuario disponer de una estimación organizada para contrastar datos con su empresa, aunque se plantea como mejora futura la exportación de informes en PDF.",
    )

    doc.add_heading("Abstract", level=1)
    add_plain_paragraph(
        doc,
        "This Final Degree Project presents the development of MyJornia, a mobile application designed for workers with variable shifts who need a clearer way to manage their work calendar, worked hours and estimated payroll. The identified technological problem is related to the fragmentation of information across calendars, personal notes and company documents, which makes it difficult to verify regular hours, overtime, night work, holidays and customized salary concepts.",
    )
    add_plain_paragraph(
        doc,
        "The main objective of the system is to develop a cross-platform tool that allows users to register shifts, configure payroll rules and consult monthly summaries through a visual mobile interface. Specific objectives were defined around interface usability, persistent data management, user authentication and the integration of security measures suitable for sensitive employment-related information.",
    )
    add_plain_paragraph(
        doc,
        "The methodology follows an MVP approach with incremental prototyping over approximately one month. The application was developed using React Native, Expo, TypeScript and Supabase, while GitHub was used for version control. The project included analysis, design, implementation, functional testing, usability validation and execution tests through Expo Go, web mode and test users.",
    )
    add_plain_paragraph(
        doc,
        "The final result is a functional application including registration and login, shift calendar, customizable shift templates, payroll settings, night work calculation, holiday management, deductions, profile configuration and security options. The solution improves the visibility of worked time and helps users compare estimated payroll information with company records. Future work includes the export of monthly reports in PDF format.",
    )


def add_intro_objetivos(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("2. Introducción y problema", level=1)
    add_plain_paragraph(
        doc,
        "La organización de turnos laborales es una necesidad habitual en sectores como hostelería, comercio, sanidad, atención al cliente y otros entornos con jornadas variables. En estos contextos, el trabajador puede acumular turnos ordinarios, turnos partidos, horas complementarias, horas nocturnas, festivos y otros conceptos que afectan a la remuneración mensual. Sin una herramienta específica, el seguimiento suele realizarse mediante notas, capturas de cuadrantes, calendarios genéricos o cálculos manuales, con el riesgo de olvidar cambios y no detectar diferencias entre lo trabajado y lo abonado.",
    )
    add_plain_paragraph(
        doc,
        "El problema abordado es relevante porque la información salarial y horaria tiene impacto directo en la economía personal del trabajador. Una estimación clara del salario bruto, pluses y deducciones permite tomar decisiones, revisar errores y disponer de una base informativa para reclamar discrepancias a la empresa. Además, la gestión visual de turnos reduce la carga cognitiva y permite consultar de forma rápida qué jornadas se han realizado, cuáles son festivas y qué periodos computan como nocturnos.",
    )
    add_plain_paragraph(
        doc,
        "MyJornia se plantea como una propuesta de mejora basada en una aplicación móvil multiplataforma. La elección de un entorno móvil responde al uso cotidiano del smartphone y a la necesidad de consultar turnos en cualquier momento. React Native y Expo permiten crear aplicaciones para Android, iOS y web a partir de un mismo proyecto JavaScript/TypeScript, lo que favorece el desarrollo incremental y la validación rápida con Expo Go (Expo, 2026a; Meta Platforms, 2026a).",
    )
    add_plain_paragraph(
        doc,
        "Se plantea que el desarrollo de una plataforma móvil de gestión de turnos y nómina permitirá optimizar el control personal de jornadas variables, mejorar la trazabilidad de los datos laborales y facilitar una estimación salarial comprensible para el usuario final. Esta hipótesis se valida mediante la implementación de las funcionalidades principales y la comprobación de su funcionamiento con usuarios de prueba.",
    )

    doc.add_heading("3. Objetivos", level=1)
    add_plain_paragraph(doc, "El objetivo general del proyecto es desarrollar una aplicación móvil multiplataforma que permita gestionar turnos laborales y estimar la nómina mensual a partir de reglas configurables por el usuario.")
    doc.add_heading("3.1 Objetivos específicos", level=2)
    add_bullets(
        doc,
        [
            "Analizar las necesidades de trabajadores con turnos variables en relación con calendario, horas, nocturnidad, festivos y salario estimado.",
            "Diseñar una interfaz móvil clara, visual y accesible que permita consultar turnos y configurar información laboral sin complejidad innecesaria.",
            "Implementar un sistema de autenticación y persistencia de datos mediante Supabase, manteniendo la separación de información por usuario.",
            "Desarrollar un calendario con vistas mensual, semanal, diaria y anual, junto con plantillas de turnos reutilizables y turnos puntuales.",
            "Configurar reglas de nómina que permitan calcular horas de contrato, horas extra, tramos nocturnos, festivos, pluses, deducciones y conceptos personalizados.",
            "Validar el funcionamiento de la aplicación mediante pruebas en Expo Go, web y cuentas de usuario de prueba.",
        ],
    )


def add_estado_arte(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("4. Estado del arte y contexto tecnológico", level=1)
    add_plain_paragraph(
        doc,
        "El estado del arte se ha analizado desde tres perspectivas: soluciones existentes de calendario y turnos, tecnologías de desarrollo móvil y fundamentos de experiencia de usuario, seguridad y arquitectura cloud. Esta revisión permite justificar por qué MyJornia combina calendario, plantillas de turnos y cálculo salarial en una misma herramienta.",
    )
    doc.add_heading("4.1 Soluciones existentes", level=2)
    add_plain_paragraph(
        doc,
        "Existen aplicaciones orientadas a calendarios laborales, registro horario y planificación de turnos. Algunas se centran en recordar horarios y otras en registrar fichajes, pero no siempre permiten adaptar tramos de nocturnidad, festivos autonómicos, pluses salariales o deducciones personalizadas. MyJornia se diferencia al priorizar la estimación individual del dinero líquido aproximado y la configuración flexible de condiciones de nómina.",
    )
    add_plain_paragraph(
        doc,
        "Frente a un calendario genérico, la aplicación propuesta incorpora conceptos propios del ámbito laboral: plantillas de turnos, turnos sin horas para vacaciones o días libres, notas, colores, iconos, horas complementarias, nocturnidad y festivos. Esta especialización permite transformar un calendario visual en una herramienta de control laboral y económico.",
    )
    doc.add_heading("4.2 Tecnologías utilizadas y comparadas", level=2)
    add_plain_paragraph(
        doc,
        "Para el desarrollo móvil se compararon alternativas como desarrollo nativo Android, Flutter y React Native. La elección de React Native con Expo se fundamenta en la posibilidad de crear una aplicación multiplataforma desde un único código base, probar rápidamente en dispositivos físicos y mantener una estructura compatible con futuras compilaciones para tiendas de aplicaciones. React Native ofrece componentes básicos como vistas, textos, imágenes y campos de entrada, adecuados para construir interfaces móviles reutilizables (Meta Platforms, 2026a).",
    )
    add_plain_paragraph(
        doc,
        "Expo se utilizó como entorno de desarrollo por simplificar la creación del proyecto, la ejecución en Android, iOS y web, y el uso de módulos nativos como almacenamiento seguro, autenticación local y pantalla de carga. La documentación oficial recomienda `create-expo-app` para iniciar proyectos React Native con menor carga de configuración inicial (Expo, 2026a).",
    )
    add_plain_paragraph(
        doc,
        "TypeScript se empleó para reforzar la calidad del código mediante tipado estático. Este enfoque reduce errores frecuentes relacionados con tipos de datos, valores nulos o estructuras inesperadas, especialmente en un proyecto con modelos de turnos, reglas de nómina, deducciones y datos de usuario (Microsoft, 2026). Node.js y npm se utilizaron como entorno de ejecución y gestión de dependencias, al tratarse de herramientas habituales en ecosistemas JavaScript modernos (OpenJS Foundation, 2026).",
    )
    add_plain_paragraph(
        doc,
        "Supabase se seleccionó como backend por integrar autenticación, base de datos PostgreSQL, API REST generada automáticamente y políticas de seguridad por filas. Supabase Auth facilita la gestión de usuarios, sesiones y recuperación de contraseña, mientras que Row Level Security permite restringir cada tabla a los datos del usuario autenticado (Supabase, 2026a, 2026b).",
    )
    doc.add_heading("4.3 Fundamentos UX/UI, cloud y seguridad", level=2)
    add_plain_paragraph(
        doc,
        "El diseño UX/UI se orientó a reducir pasos y hacer visibles los datos relevantes. Se priorizaron vistas de calendario, menús inferiores, botones reconocibles, iconos, colores diferenciados y formularios agrupados por bloques funcionales. Los criterios de accesibilidad y usabilidad se relacionan con principios de percepción, operación y comprensión de la interfaz, en línea con la orientación general de WCAG 2.2 para mejorar el acceso a contenidos digitales (W3C, 2024). Asimismo, el uso de componentes reutilizables sigue el enfoque de composición propio de React, que facilita dividir la interfaz en piezas pequeñas y mantenibles (Meta Platforms, 2026b).",
    )
    add_plain_paragraph(
        doc,
        "Desde el punto de vista de seguridad, el proyecto incorpora autenticación, almacenamiento seguro de sesión, bloqueo local opcional, doble factor y políticas RLS en la base de datos. OWASP considera la seguridad móvil como un estándar específico para aplicaciones que tratan datos sensibles, y su enfoque sirvió como referencia para integrar controles preventivos durante el desarrollo (OWASP Foundation, 2026a).",
    )
    add_plain_paragraph(
        doc,
        "Además, se consideró el principio de protección de datos desde el diseño y por defecto. La AEPD indica que los requisitos de privacidad deben incorporarse desde las primeras fases de definición y desarrollo de un sistema, evitando que la protección de datos sea una capa añadida al final del proyecto (Agencia Española de Protección de Datos, 2023, 2026). Para valorar la calidad del producto software se tuvieron presentes atributos como funcionalidad, usabilidad, fiabilidad, seguridad y mantenibilidad, relacionados con el modelo ISO/IEC 25010 (International Organization for Standardization, 2023).",
    )
    add_table(
        doc,
        "Tabla 1. Comparación tecnológica. Fuente: elaboración propia.",
        ["Tecnología", "Uso en MyJornia", "Justificación"],
        [
            ["React Native", "Interfaz móvil", "Permite construir vistas y componentes reutilizables para Android, iOS y web."],
            ["Expo", "Entorno de desarrollo", "Facilita pruebas en Expo Go, web y configuración de módulos nativos."],
            ["TypeScript", "Lenguaje principal", "Aporta tipado estático y mejora la mantenibilidad."],
            ["Supabase", "Backend y base de datos", "Integra autenticación, PostgreSQL, API y políticas RLS."],
            ["GitHub", "Control de versiones", "Permite conservar historial del código y documentación del proyecto."],
        ],
    )


def add_metodologia(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("5. Metodología de desarrollo", level=1)
    add_plain_paragraph(
        doc,
        "La metodología aplicada fue un MVP con prototipado incremental. Esta elección resultó adecuada porque el proyecto partía de una necesidad concreta, pero requería comprobar progresivamente qué pantallas y configuraciones eran más útiles para el usuario final. En lugar de cerrar todo el diseño al inicio, se desarrollaron versiones sucesivas: autenticación, calendario, plantillas de turnos, resumen, nómina, perfil y seguridad.",
    )
    add_plain_paragraph(
        doc,
        "El tiempo aproximado de desarrollo fue de un mes. Durante ese periodo se aplicaron ciclos cortos de análisis, diseño, implementación y prueba. Cada ciclo permitió detectar fallos, ajustar la interfaz, separar componentes y ampliar la funcionalidad de la aplicación sin perder el objetivo principal del sistema.",
    )
    add_table(
        doc,
        "Tabla 2. Fases de la metodología. Fuente: elaboración propia.",
        ["Fase", "Descripción", "Resultado"],
        [
            ["Análisis", "Identificación del problema, público objetivo y requisitos principales.", "Definición de funcionalidades: calendario, turnos y nómina."],
            ["Diseño", "Diseño de flujo de navegación, pantallas y modelo de datos.", "Menú inferior, vistas principales y esquema Supabase."],
            ["Desarrollo", "Implementación incremental con Expo, React Native y Supabase.", "Aplicación funcional con autenticación y persistencia."],
            ["Testing", "Pruebas en Expo Go, web y usuarios de prueba.", "Corrección de errores de inicio, guardado y navegación."],
            ["Despliegue", "Preparación de configuración, icono, splash y documentación.", "Proyecto listo para evolución hacia tiendas."],
        ],
    )
    add_plain_paragraph(
        doc,
        "La validación combinó pruebas funcionales, usabilidad, seguridad y rendimiento básico. Las pruebas funcionales verificaron el registro, inicio de sesión, recuperación de contraseña, creación de turnos, guardado de plantillas, configuración de nómina, cierre de sesión y persistencia de datos. Las pruebas de usabilidad comprobaron que el usuario pudiera localizar las ventanas principales desde el menú inferior. Las pruebas de seguridad revisaron que la sesión no dependiera de claves secretas en la app y que el acceso a la base de datos quedara protegido por políticas RLS.",
    )


def add_diseno(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("6. Diseño del sistema", level=1)
    add_plain_paragraph(
        doc,
        "El diseño del sistema se organizó en torno a cinco áreas principales: Calendario, Resumen, Nómina, Turnos y Más. Esta distribución permite separar la consulta de datos, la configuración salarial y la gestión de plantillas, evitando que una sola pantalla acumule demasiada información.",
    )
    doc.add_heading("6.1 Arquitectura general", level=2)
    add_plain_paragraph(
        doc,
        "La arquitectura adoptada puede describirse como cliente móvil con backend cloud. La aplicación React Native se ejecuta en el dispositivo del usuario y se comunica con Supabase mediante el SDK cliente. Supabase gestiona autenticación, base de datos PostgreSQL, almacenamiento de sesión, API REST y políticas de autorización. Este enfoque reduce la necesidad de desarrollar un servidor propio, pero mantiene una separación clara entre interfaz, lógica de negocio y datos.",
    )
    add_table(
        doc,
        "Tabla 3. Componentes principales de la arquitectura. Fuente: elaboración propia.",
        ["Componente", "Responsabilidad"],
        [
            ["App móvil React Native", "Mostrar pantallas, recoger datos y calcular resúmenes en la interfaz."],
            ["Expo", "Ejecutar el proyecto en desarrollo, web y dispositivos físicos."],
            ["Supabase Auth", "Gestionar registro, inicio de sesión, recuperación de contraseña y sesión."],
            ["PostgreSQL", "Guardar perfiles, turnos, plantillas, reglas de nómina y festivos."],
            ["RLS", "Asegurar que cada usuario acceda únicamente a sus propios registros."],
            ["GitHub", "Mantener historial del código fuente y facilitar la entrega del repositorio."],
        ],
    )
    add_caption(doc, "Figura 1. Arquitectura cliente-cloud de MyJornia. Fuente: elaboración propia.")
    add_bullets(
        doc,
        [
            "Cliente: aplicación React Native/Expo instalada o ejecutada en web.",
            "Comunicación: SDK de Supabase sobre API HTTPS.",
            "Backend: Supabase Auth, API y base PostgreSQL.",
            "Seguridad: autenticación, RLS, almacenamiento seguro y bloqueo local opcional.",
        ],
    )
    doc.add_heading("6.2 Modelo de base de datos", level=2)
    add_plain_paragraph(
        doc,
        "La base de datos se implementó sobre PostgreSQL mediante Supabase. Se definieron tablas para perfiles, turnos del calendario, plantillas, reglas de nómina, tramos nocturnos, festivos locales, deducciones, pagas extra y relación entre pagas y tipos de turno. PostgreSQL permite aplicar políticas de seguridad por filas, de manera que el usuario autenticado solo vea y modifique sus datos (PostgreSQL Global Development Group, 2026).",
    )
    add_table(
        doc,
        "Tabla 4. Principales tablas de la base de datos. Fuente: elaboración propia.",
        ["Tabla", "Descripción"],
        [
            ["profiles", "Datos básicos del usuario autenticado."],
            ["work_rules", "Configuración de contrato, salario base, pluses y comunidad autónoma."],
            ["shift_types", "Plantillas de turnos reutilizables con color, icono y horario."],
            ["shifts", "Turnos concretos asignados a fechas del calendario."],
            ["night_pay_ranges", "Tramos horarios de nocturnidad y precio por hora."],
            ["holidays", "Festivos locales o manuales añadidos por el usuario."],
            ["payroll_additions", "Pagas o conceptos positivos personalizados."],
            ["payroll_addition_shift_types", "Relación entre conceptos extra y plantillas donde aplican."],
            ["payroll_deductions", "Deducciones configurables por nombre y porcentaje."],
        ],
    )
    add_plain_paragraph(
        doc,
        "El enlace público del servicio de base de datos utilizado es https://rhpvhzrmwzorhaqiwhhm.supabase.co. En la memoria no se incluyen claves de API, claves secretas ni credenciales, ya que dichos valores no deben publicarse en documentos académicos ni repositorios. El acceso administrativo se conserva únicamente en el panel privado de Supabase.",
    )
    doc.add_heading("6.3 Casos de uso", level=2)
    add_table(
        doc,
        "Tabla 5. Casos de uso principales. Fuente: elaboración propia.",
        ["Caso de uso", "Actor", "Resultado esperado"],
        [
            ["Registrarse e iniciar sesión", "Usuario", "Acceso a una cuenta personal con datos separados."],
            ["Crear plantilla de turno", "Usuario", "Turno reutilizable con nombre, horario, color e icono."],
            ["Añadir turno al calendario", "Usuario", "Registro de una jornada concreta en una fecha."],
            ["Configurar nómina", "Usuario", "Definición de contrato, salario, pluses y deducciones."],
            ["Consultar resumen", "Usuario", "Visualización de horas, festivos, nocturnidad y estimación salarial."],
            ["Eliminar cuenta", "Usuario", "Borrado de usuario y datos asociados con doble confirmación."],
        ],
    )
    doc.add_heading("6.4 Diseño de interfaz", level=2)
    add_plain_paragraph(
        doc,
        "La interfaz se diseñó para parecer una herramienta móvil real, no una página informativa. El menú inferior facilita el acceso a las ventanas principales y el calendario ocupa un papel central. El uso de colores, iconos y bloques visuales permite distinguir turnos y conceptos salariales de manera rápida. Las pantallas se organizaron con títulos grandes, paneles agrupados y controles reconocibles, siguiendo criterios de coherencia visual y facilidad de lectura.",
    )
    add_figure_placeholder(doc, 2, "Pantalla de inicio de sesión de MyJornia")
    add_figure_placeholder(doc, 3, "Vista mensual del calendario con turnos")
    add_figure_placeholder(doc, 4, "Pantalla de plantillas de turnos")
    add_figure_placeholder(doc, 5, "Pantalla de nómina y configuración salarial")
    add_figure_placeholder(doc, 6, "Pantalla de perfil, preferencias y seguridad")


def add_desarrollo(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("7. Desarrollo e implementación", level=1)
    doc.add_heading("7.1 Estructura del proyecto", level=2)
    add_plain_paragraph(
        doc,
        "El proyecto se organizó separando pantallas, componentes, hooks, utilidades y configuración de Supabase. Esta separación mejoró la mantenibilidad respecto a una implementación concentrada en un único archivo. Las pantallas principales se ubican en `src/screens`, los componentes reutilizables en `src/components`, la lógica de estado y acceso a datos en `src/hooks`, y la conexión con Supabase en `lib/supabase.ts`.",
    )
    add_table(
        doc,
        "Tabla 6. Organización del código. Fuente: elaboración propia.",
        ["Carpeta o archivo", "Función"],
        [
            ["App.tsx", "Entrada principal y composición de pantallas."],
            ["src/screens", "Pantallas de autenticación, calendario, nómina, turnos, resumen y perfil."],
            ["src/components", "Componentes reutilizables como navegación, inputs, selectores y tarjetas."],
            ["src/hooks", "Controladores de autenticación, datos, MFA y seguridad."],
            ["src/utils", "Funciones de fechas, cálculo de horas y nómina."],
            ["lib/supabase.ts", "Configuración del cliente Supabase."],
            ["docs", "Documentación legal y páginas para GitHub Pages."],
        ],
    )
    doc.add_heading("7.2 Autenticación y seguridad", level=2)
    add_plain_paragraph(
        doc,
        "El sistema de autenticación permite registro con correo y contraseña, inicio de sesión, recuperación de contraseña y cierre de sesión. La contraseña requiere una longitud mínima, un número y un símbolo. Cuando se intenta registrar un correo ya utilizado, la aplicación muestra un mensaje claro orientando al usuario a iniciar sesión o recuperar la contraseña, sin revelar información sensible de forma innecesaria.",
    )
    add_plain_paragraph(
        doc,
        "La seguridad se reforzó mediante almacenamiento seguro de sesión con Expo SecureStore, bloqueo local opcional con autenticación biométrica o código del dispositivo, doble factor y políticas RLS en Supabase. SecureStore permite guardar pares clave-valor cifrados localmente, mientras que LocalAuthentication permite integrar huella o reconocimiento facial según capacidades del dispositivo (Expo, 2026b, 2026c).",
    )
    add_plain_paragraph(
        doc,
        "Además, se implementó la eliminación real de cuenta mediante una función SQL en Supabase que elimina los datos del usuario autenticado y finalmente su registro en `auth.users`. La app nunca utiliza claves `service_role`, ya que estas claves no deben exponerse en clientes móviles. Esta decisión se alinea con la recomendación de no incluir credenciales privilegiadas en aplicaciones distribuidas al usuario final.",
    )
    doc.add_heading("7.3 Calendario y turnos", level=2)
    add_plain_paragraph(
        doc,
        "El calendario permite consultar turnos por mes, semana, día y año. El usuario puede cambiar de periodo y seleccionar días para añadir turnos. Se implementaron dos flujos: crear un turno puntual, que no se guarda como plantilla, y añadir un turno guardado desde la ventana de plantillas. Esta separación responde a un uso real: algunos turnos se repiten y otros son excepcionales.",
    )
    add_plain_paragraph(
        doc,
        "Cada turno puede incluir nombre, color, icono, horario, descanso y notas. También se contemplan turnos sin horas, útiles para vacaciones, días libres o situaciones que deben verse en el calendario pero no computar como horas trabajadas. El cálculo de horas descuenta descansos y controla turnos que cruzan medianoche.",
    )
    doc.add_heading("7.4 Nómina, resumen y cálculo salarial", level=2)
    add_plain_paragraph(
        doc,
        "La ventana de nómina permite configurar horas de contrato mensuales, salario bruto base, precio de horas complementarias, festivos, tramos nocturnos, deducciones y conceptos positivos personalizados. Los campos admiten números decimales, incluyendo coma decimal, para adaptarse a importes reales como 1,546 euros por hora.",
    )
    add_plain_paragraph(
        doc,
        "Los tramos nocturnos pueden definirse de forma independiente. Por ejemplo, se puede configurar un tramo de 22:00 a 00:00 y otro de 00:00 a 06:00 con importes diferentes. También se incorpora la comunidad autónoma para cargar festivos nacionales y autonómicos, junto con festivos locales introducidos manualmente.",
    )
    add_plain_paragraph(
        doc,
        "El resumen calcula horas trabajadas, horas de contrato, horas complementarias, nocturnidad, festivos, salario bruto estimado, pluses, deducciones y salario neto estimado. La finalidad no es sustituir una nómina oficial, sino ofrecer al trabajador una referencia organizada para revisar discrepancias y conservar trazabilidad personal.",
    )
    doc.add_heading("7.5 Repositorio y base de datos", level=2)
    add_plain_paragraph(
        doc,
        "El código fuente se gestiona mediante GitHub, que actúa como repositorio remoto y conserva el historial de cambios del proyecto. GitHub define un repositorio como el lugar donde se almacena el código, los archivos y el historial de revisiones de un proyecto (GitHub, 2026).",
    )
    add_bullets(
        doc,
        [
            "Repositorio del proyecto: https://github.com/jairopanait/MyJornia",
            "Proveedor de base de datos: Supabase PostgreSQL.",
            "URL pública del servicio Supabase: https://rhpvhzrmwzorhaqiwhhm.supabase.co",
            "Documentación legal publicada mediante GitHub Pages: https://jairopanait.github.io/MyJornia/",
        ],
    )


def add_resultados(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("8. Resultados y validación", level=1)
    add_plain_paragraph(
        doc,
        "El resultado obtenido es una aplicación funcional que cubre las ventanas principales definidas en los objetivos: autenticación, calendario, resumen, nómina, plantillas de turnos, perfil, preferencias, ayuda, seguridad y eliminación de cuenta. La aplicación se ejecutó en Expo Go, navegador web y cuentas de usuario de prueba.",
    )
    add_table(
        doc,
        "Tabla 7. Validación de funcionalidades. Fuente: elaboración propia.",
        ["Funcionalidad", "Prueba realizada", "Resultado"],
        [
            ["Registro e inicio de sesión", "Alta de usuario, login y mensajes de error.", "Correcto."],
            ["Recuperación de contraseña", "Envío de correo y apertura mediante URL configurada.", "Correcto."],
            ["Calendario", "Cambio de vista, selección de día y carga de turnos.", "Correcto."],
            ["Plantillas de turnos", "Creación, edición, color, icono y turnos sin horas.", "Correcto."],
            ["Nómina", "Guardado de reglas, decimales, nocturnidad, festivos y deducciones.", "Correcto."],
            ["Persistencia", "Cierre de sesión y nueva entrada con datos guardados.", "Correcto tras ajuste de inicialización."],
            ["Seguridad", "Bloqueo local, almacenamiento seguro y RLS.", "Correcto en entorno de pruebas."],
            ["Eliminación de cuenta", "Botón con doble confirmación y función SQL segura.", "Preparado para ejecución en Supabase."],
        ],
    )
    doc.add_heading("8.1 Resultados positivos", level=2)
    add_bullets(
        doc,
        [
            "Se implementó una navegación móvil clara mediante menú inferior.",
            "Se creó un calendario visual con turnos coloreados, iconos y notas.",
            "Se desarrolló una ventana de plantillas que permite reutilizar turnos habituales.",
            "Se integró una configuración de nómina flexible para contrato, tramos nocturnos, festivos, pagas extra y deducciones.",
            "Se añadió autenticación, recuperación de contraseña, almacenamiento seguro y bloqueo local opcional.",
            "Se documentaron políticas de privacidad y eliminación de cuenta para futura publicación.",
        ],
    )
    doc.add_heading("8.2 Resultados limitantes", level=2)
    add_plain_paragraph(
        doc,
        "Las principales limitaciones se relacionan con el alcance temporal del proyecto. Aunque la aplicación es funcional, todavía no se ha publicado en App Store o Play Store, no se ha implementado exportación en PDF y la validación se realizó con usuarios de prueba, no con un estudio amplio de usuarios reales. También queda pendiente una batería de pruebas automatizadas más extensa.",
    )
    add_plain_paragraph(
        doc,
        "En términos de escalabilidad, Supabase permite evolucionar el proyecto, pero sería necesario revisar índices, límites de uso, copias de seguridad, monitorización y costes antes de una explotación pública con muchos usuarios. La seguridad básica está planteada, aunque una versión comercial debería incorporar auditorías adicionales siguiendo estándares como OWASP ASVS y MASVS (OWASP Foundation, 2026a, 2026b).",
    )
    add_table(
        doc,
        "Tabla 8. Comparación entre objetivos y resultados. Fuente: elaboración propia.",
        ["Objetivo previsto", "Resultado real"],
        [
            ["Gestionar turnos en calendario", "Cumplido mediante vistas y turnos personalizables."],
            ["Crear plantillas reutilizables", "Cumplido con pantalla de turnos y selección desde calendario."],
            ["Calcular nómina estimada", "Cumplido con horas, nocturnidad, festivos, pluses y deducciones."],
            ["Guardar datos por usuario", "Cumplido mediante Supabase y políticas RLS."],
            ["Proteger información sensible", "Cumplido de forma inicial con autenticación, RLS, SecureStore y bloqueo."],
            ["Exportar datos en PDF", "Pendiente como mejora futura."],
        ],
    )


def add_discusion_conclusiones(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("9. Discusión", level=1)
    add_plain_paragraph(
        doc,
        "Los resultados permiten confirmar que una aplicación móvil específica puede mejorar la organización personal de turnos y la estimación salarial de trabajadores con horarios variables. Frente a soluciones genéricas, MyJornia integra calendario y nómina en un mismo flujo, reduciendo la necesidad de trasladar datos entre herramientas separadas.",
    )
    add_plain_paragraph(
        doc,
        "La elección de React Native y Expo facilitó el desarrollo en un periodo breve, especialmente por la posibilidad de probar en Expo Go y web. Sin embargo, algunas funciones de seguridad nativa presentan diferencias entre Expo Go y una compilación final, como ocurre con ciertas limitaciones de autenticación biométrica en iOS indicadas por la documentación de Expo (Expo, 2026c). Por ello, una futura versión de producción debería validarse mediante builds reales.",
    )
    add_plain_paragraph(
        doc,
        "La base de datos cloud con Supabase permitió avanzar con rapidez sin crear un backend propio. No obstante, esta decisión implica depender de un proveedor externo y exige comprender correctamente RLS para evitar exposiciones de datos. La documentación de Supabase y PostgreSQL destaca que las políticas por filas son esenciales cuando las tablas expuestas se consultan desde clientes autenticados (Supabase, 2026b; PostgreSQL Global Development Group, 2026).",
    )
    doc.add_heading("9.1 Limitaciones", level=2)
    add_bullets(
        doc,
        [
            "La aplicación no está publicada todavía en tiendas oficiales.",
            "La exportación de informes en PDF queda pendiente.",
            "La validación se realizó con pruebas funcionales y usuarios de prueba, no con un estudio estadístico amplio.",
            "El cálculo salarial es una estimación y puede variar según convenios, retenciones reales y reglas empresariales.",
            "Las políticas de seguridad deberían auditarse antes de una explotación pública.",
        ],
    )
    doc.add_heading("9.2 Mejoras futuras", level=2)
    add_bullets(
        doc,
        [
            "Exportar resúmenes mensuales en PDF.",
            "Añadir widgets para consultar próximos turnos desde la pantalla del móvil.",
            "Incorporar notificaciones y recordatorios de turnos.",
            "Permitir copias de seguridad o sincronización avanzada.",
            "Añadir pruebas automatizadas y monitorización de errores.",
            "Preparar builds de producción para App Store y Play Store.",
        ],
    )
    doc.add_heading("10. Conclusiones", level=1)
    add_plain_paragraph(
        doc,
        "El proyecto MyJornia demuestra la viabilidad de desarrollar una aplicación móvil multiplataforma orientada al control personal de turnos y nómina en un periodo aproximado de un mes. Se alcanzó una versión funcional con autenticación, calendario, plantillas, configuración salarial, resumen mensual, perfil y medidas de seguridad iniciales.",
    )
    add_plain_paragraph(
        doc,
        "El objetivo general se considera cumplido, ya que la aplicación permite registrar jornadas, configurar reglas laborales y consultar una estimación mensual organizada. Los objetivos específicos también se alcanzaron en su mayoría, salvo la exportación de PDF, que se define como mejora futura por limitación temporal.",
    )
    add_plain_paragraph(
        doc,
        "La principal aportación del trabajo es combinar en una misma herramienta elementos que normalmente se encuentran separados: calendario laboral, turnos personalizados, cálculo de horas y aproximación salarial. Esta integración aporta valor a trabajadores de hostelería, comercio y otros sectores con turnos variables, al facilitar una visión más transparente de su tiempo trabajado y del dinero que podrían percibir.",
    )
    add_plain_paragraph(
        doc,
        "Desde una perspectiva académica, el proyecto permitió aplicar conocimientos de desarrollo multiplataforma, bases de datos, APIs, autenticación, ciberseguridad, diseño de interfaz y control de versiones. Además, se reforzó la importancia de diseñar desde el inicio con criterios de privacidad, mantenibilidad y evolución futura.",
    )


def add_bibliografia(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("11. Bibliografía", level=1)
    refs = [
        "Agencia Española de Protección de Datos. (2023). Protección de datos desde el diseño. https://www.aepd.es/derechos-y-deberes/cumple-tus-deberes/medidas-de-cumplimiento/proteccion-de-datos-desde-el-diseno",
        "Agencia Española de Protección de Datos. (2026). Protección de datos por defecto. https://www.aepd.es/en/rights-and-duties/fulfill-your-duties/measures-compliance/privacy-default",
        "Expo. (2026a). Create a project. Expo Documentation. https://docs.expo.dev/get-started/create-a-project/",
        "Expo. (2026b). SecureStore. Expo Documentation. https://docs.expo.dev/versions/latest/sdk/securestore/",
        "Expo. (2026c). LocalAuthentication. Expo Documentation. https://docs.expo.dev/versions/latest/sdk/local-authentication/",
        "GitHub. (2026). About repositories. GitHub Docs. https://docs.github.com/en/repositories/creating-and-managing-repositories/about-repositories",
        "International Organization for Standardization. (2023). ISO/IEC 25010:2023 Systems and software engineering - Product quality model. https://www.iso.org/standard/78176.html",
        "Meta Platforms. (2026a). Core Components and APIs. React Native. https://reactnative.dev/docs/components-and-apis",
        "Meta Platforms. (2026b). Quick Start. React. https://react.dev/learn",
        "Microsoft. (2026). The TypeScript Handbook. https://www.typescriptlang.org/docs/handbook/",
        "OpenJS Foundation. (2026). About Node.js. https://nodejs.org/en/about/",
        "OWASP Foundation. (2026a). OWASP Mobile Application Security. https://owasp.org/www-project-mobile-app-security/",
        "OWASP Foundation. (2026b). OWASP Application Security Verification Standard. https://owasp.org/www-project-application-security-verification-standard/",
        "PostgreSQL Global Development Group. (2026). Row Security Policies. PostgreSQL 17 Documentation. https://www.postgresql.org/docs/17/ddl-rowsecurity.html",
        "Supabase. (2026a). Auth. Supabase Docs. https://supabase.com/docs/guides/auth",
        "Supabase. (2026b). Row Level Security. Supabase Docs. https://supabase.com/docs/guides/database/postgres/row-level-security",
        "World Wide Web Consortium. (2024). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/wcag/",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_anexos(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("12. Anexos", level=1)
    doc.add_heading("Anexo A. Checklist de cumplimiento de la guía", level=2)
    add_table(
        doc,
        "Tabla 9. Lista de comprobación de requisitos de memoria. Fuente: elaboración propia.",
        ["Requisito", "Cumplimiento en la memoria"],
        [
            ["Portada", "Incluye título, tema, autor, centro, ciclo, tutor/a y fecha."],
            ["Índice", "Se incluye índice automático actualizable en Word."],
            ["Resumen y abstract", "Incluidos en español e inglés con cuatro bloques temáticos."],
            ["Objetivos", "Un objetivo general y varios objetivos específicos redactados en infinitivo."],
            ["Marco teórico", "Incluye soluciones, tecnologías y fundamentos UX/UI, cloud y seguridad."],
            ["Metodología", "Describe MVP, prototipado, fases, testing y validación."],
            ["Resultados", "Compara objetivos previstos y resultados reales."],
            ["Conclusiones", "Incluye discusión, limitaciones, aplicabilidad y mejoras futuras."],
            ["APA", "Incluye citas en el texto y bibliografía con más de 15 fuentes."],
            ["Figuras y tablas", "Se numeran y se explican dentro del documento."],
        ],
    )
    doc.add_heading("Anexo B. Capturas pendientes de insertar", level=2)
    add_plain_paragraph(
        doc,
        "Cuando se disponga de capturas reales, deberán sustituirse los marcadores de figuras por imágenes de la aplicación. Se recomienda incluir capturas de inicio de sesión, calendario mensual, plantillas de turnos, configuración de nómina, resumen y perfil.",
    )
    add_numbered(
        doc,
        [
            "Inicio de sesión y registro.",
            "Calendario mensual con turnos.",
            "Creación o edición de turno.",
            "Pantalla de plantillas de turnos.",
            "Pantalla de nómina y tramos nocturnos.",
            "Resumen mensual de horas y salario.",
            "Perfil, preferencias, seguridad y eliminación de cuenta.",
        ],
    )
    doc.add_heading("Anexo C. Enlaces del proyecto", level=2)
    add_bullets(
        doc,
        [
            "Repositorio GitHub: https://github.com/jairopanait/MyJornia",
            "Base de datos: Supabase PostgreSQL con acceso administrativo privado.",
            "URL pública Supabase: https://rhpvhzrmwzorhaqiwhhm.supabase.co",
            "Política de privacidad: https://jairopanait.github.io/MyJornia/privacidad/",
            "Eliminación de cuenta: https://jairopanait.github.io/MyJornia/eliminar-cuenta/",
        ],
    )


def main() -> None:
    doc = setup_document()
    cover(doc)
    add_front_matter(doc)
    add_resumen(doc)
    add_intro_objetivos(doc)
    add_estado_arte(doc)
    add_metodologia(doc)
    add_diseno(doc)
    add_desarrollo(doc)
    add_resultados(doc)
    add_discusion_conclusiones(doc)
    add_bibliografia(doc)
    add_anexos(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
